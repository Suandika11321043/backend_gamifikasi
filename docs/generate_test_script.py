"""Generate Test Script Word document — format satu tabel utuh per butir uji (sesuai template)."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"d:\TA2\gamifikasi\docs\Test_Script_Gamifikasi_PAUD.docx"

TESTER = "_________________________"
TEST_DATE = "_________________________"
FONT = "Times New Roman"
COLS = 4  # tabel dasar 4 kolom; metadata pakai label + merge 3 kolom kanan

TEST_CASES = [
    {
        "id": "BU-01",
        "nama": "Login admin ke sistem",
        "tujuan": "Memastikan admin dapat masuk ke sistem dengan kredensial yang valid",
        "deskripsi": "Fitur ini digunakan oleh admin untuk mengakses halaman dashboard dan seluruh modul manajemen.",
        "kondisi_awal": (
            "1. Backend Spring Boot berjalan di http://localhost:8081\n"
            "2. Frontend React berjalan (Vite dev server)\n"
            "3. Akun admin sudah terdaftar di database"
        ),
        "skenario": [
            "Buka halaman login aplikasi",
            "Masukkan username dan password admin yang valid",
            "Klik tombol Login",
        ],
        "kriteria": (
            "Pengujian berhasil apabila admin berhasil masuk, token tersimpan, "
            "dan diarahkan ke halaman dashboard admin."
        ),
        "kasus": [
            {
                "masukan": "Masukkan username dan password admin yang valid, lalu klik Login",
                "diharapkan": "Admin berhasil masuk dan diarahkan ke halaman dashboard",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-02",
        "nama": "Menambah tema pembelajaran",
        "tujuan": "Memastikan admin dapat menambah tema baru beserta ikon",
        "deskripsi": "Fitur ini digunakan admin pada halaman Manajemen Tema untuk menambah topik pembelajaran PAUD.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Admin berada di halaman Manajemen Tema (/admin/tema)"
        ),
        "skenario": [
            "Klik tombol Tambah Tema",
            "Isi nama tema dan deskripsi",
            "Pilih file ikon gambar (JPG/PNG, maks. 5MB)",
            "Atur status tema (Aktif/Nonaktif)",
            "Klik Simpan",
        ],
        "kriteria": (
            "Pengujian berhasil apabila tema baru muncul di daftar tema, "
            "ikon tampil di kartu tema, dan data tersimpan di database."
        ),
        "kasus": [
            {
                "masukan": "Isi nama tema, unggah ikon valid, klik Simpan",
                "diharapkan": "Tema baru tersimpan dan muncul di daftar tema",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-03",
        "nama": "Mengedit tema dan preview ikon",
        "tujuan": "Memastikan admin dapat mengubah data tema dan melihat preview ikon baru",
        "deskripsi": "Fitur edit tema memungkinkan admin memperbarui nama, deskripsi, status, dan ikon tema.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Minimal satu tema sudah ada di sistem\n"
            "3. Admin berada di halaman Manajemen Tema"
        ),
        "skenario": [
            "Klik ikon edit (pensil) pada salah satu kartu tema",
            "Ubah nama atau deskripsi tema",
            "Pilih file ikon baru",
            "Periksa preview ikon di modal sebelum menyimpan",
            "Klik Simpan",
        ],
        "kriteria": (
            "Pengujian berhasil apabila preview ikon baru tampil sebelum simpan, "
            "data tema terupdate, dan ikon baru tampil di daftar tema."
        ),
        "kasus": [
            {
                "masukan": "Ganti ikon tema dengan file gambar baru, periksa preview, lalu klik Simpan",
                "diharapkan": "Preview ikon baru tampil di modal dan ikon terbaru muncul di kartu tema",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-04",
        "nama": "Menghapus tema beserta data terkait",
        "tujuan": "Memastikan admin dapat menghapus tema tanpa error foreign key",
        "deskripsi": (
            "Fitur ini digunakan admin untuk menghapus tema secara permanen. "
            "Sistem menghapus cascade: soal, opsi, jawaban siswa, skor, dan sesi timer terkait."
        ),
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Tema yang akan dihapus memiliki soal dan data terkait\n"
            "3. Admin berada di halaman Manajemen Tema"
        ),
        "skenario": [
            "Pilih tema yang memiliki soal terkait",
            "Klik ikon hapus (tempat sampah) pada kartu tema",
            "Baca konfirmasi di modal Hapus Tema?",
            "Klik tombol Hapus",
        ],
        "kriteria": (
            "Pengujian berhasil apabila tema terhapus dari daftar, "
            "tidak muncul error FK constraint, dan data terkait (soal, skor) ikut terhapus."
        ),
        "kasus": [
            {
                "masukan": "Klik tombol hapus pada tema yang memiliki soal, konfirmasi dengan klik Hapus",
                "diharapkan": "Tema terhapus dari daftar dan tidak muncul error database",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-05",
        "nama": "Mengaktifkan/nonaktifkan tema",
        "tujuan": "Memastikan admin dapat mengubah status ketersediaan tema",
        "deskripsi": "Toggle status tema mengontrol apakah tema dapat diakses siswa.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Minimal satu tema ada di halaman Manajemen Tema"
        ),
        "skenario": [
            "Pada kartu tema, klik toggle status Aktif/Nonaktif",
            "Amati perubahan label status pada kartu",
            "Buka halaman siswa dan periksa daftar tema",
        ],
        "kriteria": (
            "Pengujian berhasil apabila status tema berubah di admin dan "
            "tema nonaktif tidak dapat diakses siswa."
        ),
        "kasus": [
            {
                "masukan": "Nonaktifkan tema yang sedang Aktif melalui toggle pada kartu tema",
                "diharapkan": "Status berubah menjadi Nonaktif dan tema tidak tampil untuk siswa",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-06",
        "nama": "Menambah dan mengedit data siswa",
        "tujuan": "Memastikan admin dapat mengelola data siswa PAUD",
        "deskripsi": "Fitur manajemen siswa untuk menambah, mengedit, dan menghapus profil siswa beserta avatar.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Admin berada di halaman Daftar Siswa (/admin/siswa)"
        ),
        "skenario": [
            "Klik Tambah Siswa",
            "Isi nama dan kelompok siswa",
            "Upload foto avatar (opsional)",
            "Klik Simpan",
        ],
        "kriteria": (
            "Pengujian berhasil apabila data siswa tersimpan, avatar tampil, "
            "dan validasi menolak input kosong."
        ),
        "kasus": [
            {
                "masukan": "Isi nama dan kelompok siswa, klik Simpan",
                "diharapkan": "Siswa baru muncul di tabel daftar siswa",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-07",
        "nama": "Menghapus data siswa",
        "tujuan": "Memastikan admin dapat menghapus siswa beserta data terkait",
        "deskripsi": "Penghapusan siswa menghapus jawaban, skor, dan sesi timer siswa tersebut.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Siswa yang akan dihapus sudah terdaftar\n"
            "3. Admin berada di halaman Daftar Siswa"
        ),
        "skenario": [
            "Pilih siswa pada tabel",
            "Klik ikon hapus",
            "Konfirmasi hapus pada modal dengan menekan tombol Hapus",
        ],
        "kriteria": (
            "Pengujian berhasil apabila siswa terhapus dari daftar "
            "dan tidak muncul lagi di halaman daftar siswa."
        ),
        "kasus": [
            {
                "masukan": "Klik tombol hapus pada salah satu siswa, konfirmasi dengan klik Hapus",
                "diharapkan": "Data siswa terhapus dari database dan tidak muncul lagi di halaman daftar",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-08",
        "nama": "Menambah soal pembelajaran",
        "tujuan": "Memastikan admin dapat menambah soal pada tema dan tanggal belajar",
        "deskripsi": (
            "Fitur manajemen soal untuk membuat soal tipe QUIZ, SORTING, MATCH, "
            "DRAG_AND_DROP, dan PUZZLE pada tanggal belajar tertentu."
        ),
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Tema aktif sudah ada\n"
            "3. Admin berada di halaman soal per tanggal (/admin/soal/{topicId}/date/{date})"
        ),
        "skenario": [
            "Klik Tambah Soal",
            "Pilih tipe soal (mis. QUIZ)",
            "Isi instruksi soal dan poin skor (minimal 1)",
            "Opsional: unggah gambar/audio, atur batas waktu",
            "Klik Simpan",
        ],
        "kriteria": (
            "Pengujian berhasil apabila soal tersimpan dan muncul di daftar soal hari tersebut."
        ),
        "kasus": [
            {
                "masukan": "Isi instruksi soal, poin skor, pilih tipe QUIZ, klik Simpan",
                "diharapkan": "Soal baru tersimpan dan tampil di daftar soal tanggal tersebut",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-09",
        "nama": "Mengatur ketersediaan soal per tanggal",
        "tujuan": "Memastikan admin dapat membuka/mengunci soal pada tanggal belajar",
        "deskripsi": (
            "Fitur kunci/buka soal per tanggal agar siswa hanya dapat mengerjakan "
            "soal yang sudah diaktifkan admin."
        ),
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Tema memiliki soal pada minimal satu tanggal belajar\n"
            "3. Admin berada di halaman kalender soal per tema"
        ),
        "skenario": [
            "Pada daftar tanggal, klik toggle kunci/buka (Lock/Unlock)",
            "Periksa indikator status ketersediaan tanggal",
            "Login sebagai siswa dan coba akses tanggal tersebut",
        ],
        "kriteria": (
            "Pengujian berhasil apabila status ketersediaan berubah dan "
            "siswa hanya dapat mengerjakan soal pada tanggal yang dibuka."
        ),
        "kasus": [
            {
                "masukan": "Klik toggle untuk membuka soal pada tanggal yang terkunci",
                "diharapkan": "Status berubah terbuka dan siswa dapat mengakses soal pada tanggal tersebut",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-10",
        "nama": "Reschedule (pindah tanggal) soal",
        "tujuan": "Memastikan admin dapat memindahkan semua soal ke tanggal belajar lain",
        "deskripsi": "Fitur reschedule memindahkan seluruh soal dari satu tanggal ke tanggal baru.",
        "kondisi_awal": (
            "1. Admin sudah login\n"
            "2. Tema memiliki soal pada tanggal A\n"
            "3. Tanggal B belum memiliki soal dan bukan Sabtu/Minggu"
        ),
        "skenario": [
            "Buka halaman soal tanggal A atau kalender tema",
            "Klik Reschedule / Pindah Tanggal",
            "Pilih tanggal B (bukan akhir pekan)",
            "Klik Simpan",
        ],
        "kriteria": (
            "Pengujian berhasil apabila semua soal pindah ke tanggal B "
            "dan tanggal lama tidak lagi memiliki soal."
        ),
        "kasus": [
            {
                "masukan": "Pindahkan soal ke tanggal Senin–Jumat yang belum memiliki soal",
                "diharapkan": "Semua soal muncul di tanggal baru dan tanggal lama kosong",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-11",
        "nama": "Siswa mengerjakan quiz",
        "tujuan": "Memastikan siswa dapat mengerjakan soal pada tema dan tanggal yang tersedia",
        "deskripsi": (
            "Alur siswa: pilih profil → pilih tema → pilih minggu/tanggal → "
            "kerjakan soal → terima feedback benar/salah."
        ),
        "kondisi_awal": (
            "1. Backend dan frontend berjalan\n"
            "2. Tema aktif dengan soal tersedia (unlocked) pada tanggal uji\n"
            "3. Siswa sudah terdaftar"
        ),
        "skenario": [
            "Buka halaman landing siswa",
            "Pilih profil siswa",
            "Pilih tema pembelajaran",
            "Pilih tanggal/minggu yang tersedia",
            "Jawab soal dan submit",
        ],
        "kriteria": (
            "Pengujian berhasil apabila soal tampil, jawaban dinilai, "
            "dan feedback popup benar/salah muncul."
        ),
        "kasus": [
            {
                "masukan": "Jawab soal QUIZ dengan opsi yang benar, lalu submit",
                "diharapkan": "Feedback benar ditampilkan dan skor/bintang bertambah",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-12",
        "nama": "Timer quiz dan penyimpanan progres",
        "tujuan": "Memastikan timer soal berjalan dan progres disimpan ke server",
        "deskripsi": (
            "Timer countdown ditampilkan di UI siswa. "
            "Progres timer dikirim ke POST /api/quiz/timer setiap 20 detik."
        ),
        "kondisi_awal": (
            "1. Soal memiliki batas waktu (timeLimitMinutes)\n"
            "2. Siswa sedang mengerjakan quiz"
        ),
        "skenario": [
            "Buka soal dengan batas waktu",
            "Amati countdown timer di layar",
            "Tunggu minimal 20 detik tanpa submit",
            "Lanjutkan atau refresh halaman quiz",
        ],
        "kriteria": (
            "Pengujian berhasil apabila timer countdown berjalan "
            "dan request timer terkirim ke server secara berkala."
        ),
        "kasus": [
            {
                "masukan": "Kerjakan soal ber-timer selama lebih dari 20 detik",
                "diharapkan": "Countdown berjalan dan POST /api/quiz/timer terpanggil",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
    {
        "id": "BU-13",
        "nama": "Riwayat jawaban siswa",
        "tujuan": "Memastikan siswa dapat melihat riwayat jawaban yang pernah dikerjakan",
        "deskripsi": "Halaman history menampilkan jawaban siswa per topik dan tanggal.",
        "kondisi_awal": (
            "1. Siswa sudah pernah mengerjakan minimal satu soal\n"
            "2. Siswa berada di halaman riwayat jawaban"
        ),
        "skenario": [
            "Pilih profil siswa yang memiliki riwayat",
            "Buka menu Riwayat Jawaban",
            "Scroll daftar riwayat",
            "Buka detail salah satu jawaban",
        ],
        "kriteria": (
            "Pengujian berhasil apabila riwayat jawaban tampil lengkap "
            "dengan informasi soal, tanggal, dan status benar/salah."
        ),
        "kasus": [
            {
                "masukan": "Buka halaman riwayat jawaban siswa yang sudah pernah mengerjakan soal",
                "diharapkan": "Daftar riwayat tampil dengan data jawaban yang benar",
                "pengamatan": "",
                "kesimpulan": "[ ] diterima",
            },
        ],
    },
]


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def merge_row_span(table, row_idx, col_start, col_end):
    """Merge cells horizontally from col_start to col_end (inclusive)."""
    row = table.rows[row_idx]
    start = row.cells[col_start]
    for c in range(col_start + 1, col_end + 1):
        start.merge(row.cells[c])
    return start


def write_cell(cell, text, *, bold=False, center=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    set_cell_margins(cell)


def write_numbered_cell(cell, items, size=11):
    cell.text = ""
    for i, item in enumerate(items, 1):
        p = cell.paragraphs[0] if i == 1 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{i}.\t{item}")
        run.font.name = FONT
        run.font.size = Pt(size)
    set_cell_margins(cell)


def add_unified_test_table(doc, tc):
    """Satu tabel besar per butir uji — format persis seperti template."""
    n_kasus = len(tc["kasus"])
    # 7 meta + 2 skenario + 2 kriteria + 2 kasus header + 1 sub-header + n_kasus data
    total_rows = 7 + 2 + 2 + 2 + 1 + n_kasus
    table = doc.add_table(rows=total_rows, cols=COLS)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = 0
    meta = [
        ("Identifikasi", tc["id"]),
        ("Nama Butir Uji", tc["nama"]),
        ("Tujuan", tc["tujuan"]),
        ("Deskripsi", tc["deskripsi"]),
        ("Kondisi Awal", tc["kondisi_awal"]),
        ("Tanggal Pengujian", TEST_DATE),
        ("Penguji", TESTER),
    ]
    for label, value in meta:
        write_cell(table.rows[row].cells[0], label, bold=True)
        set_cell_shading(table.rows[row].cells[0], "FFFFFF")
        merged = merge_row_span(table, row, 1, COLS - 1)
        write_cell(merged, value)
        row += 1

    # Skenario pengujian
    merged = merge_row_span(table, row, 0, COLS - 1)
    write_cell(merged, "Skenario pengujian", bold=True, center=True)
    row += 1
    merged = merge_row_span(table, row, 0, COLS - 1)
    write_numbered_cell(merged, tc["skenario"])
    row += 1

    # Kriteria Evaluasi Hasil
    merged = merge_row_span(table, row, 0, COLS - 1)
    write_cell(merged, "Kriteria Evaluasi Hasil", bold=True, center=True)
    row += 1
    merged = merge_row_span(table, row, 0, COLS - 1)
    write_cell(merged, tc["kriteria"])
    row += 1

    # Kasus dan Hasil Pengujian — header section
    merged = merge_row_span(table, row, 0, COLS - 1)
    write_cell(merged, "Kasus dan Hasil Pengujian", bold=True, center=True)
    row += 1

    # Sub-header 4 kolom
    headers = ["Data Masukan", "Yang Diharapkan", "Pengamatan", "Kesimpulan"]
    hdr_row = table.rows[row]
    for j, h in enumerate(headers):
        write_cell(hdr_row.cells[j], h, bold=True, center=True, size=10)
        set_cell_shading(hdr_row.cells[j], "FFFFFF")
    row += 1

    # Baris data kasus
    for k in tc["kasus"]:
        data_row = table.rows[row]
        write_cell(data_row.cells[0], k["masukan"], size=10)
        write_cell(data_row.cells[1], k["diharapkan"], size=10)
        write_cell(data_row.cells[2], k.get("pengamatan", ""), size=10)
        write_cell(data_row.cells[3], k.get("kesimpulan", "[ ] diterima"), size=10)
        row += 1

    # Lebar kolom proporsional
    for tbl_row in table.rows:
        tbl_row.cells[0].width = Cm(3.2)
        for j in range(1, COLS):
            tbl_row.cells[j].width = Cm(4.3)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    for idx, tc in enumerate(TEST_CASES):
        add_unified_test_table(doc, tc)
        if idx < len(TEST_CASES) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_document()
