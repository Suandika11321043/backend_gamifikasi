#!/usr/bin/env python3
"""Generate Use Case Specification document (.docx) for Gamifikasi PAUD."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "Use_Case_Specification_Gamifikasi_PAUD_v3.docx"

USE_CASES = [
    # ── UC-01 Login ──────────────────────────────────────────────
    {
        "id": "UC-01",
        "name": "Login Admin",
        "description": (
            "Use case ini menggambarkan proses login yang dilakukan oleh admin (guru) "
            "untuk mengakses halaman dashboard pada Sistem Gamifikasi Pembelajaran PAUD. "
            "Proses autentikasi mencakup validasi username dan kata sandi sebelum admin "
            "diberikan hak akses ke dashboard."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin memiliki akun yang valid dan aktif pada sistem.",
        "post": "Admin berhasil login dan diarahkan ke halaman dashboard sesuai hak akses.",
        "basic_flow": [
            ("Admin membuka halaman login", "Sistem menampilkan form login berisi field username dan kata sandi"),
            ("Admin mengisi username dan kata sandi yang valid", "Sistem memvalidasi kredensial yang dimasukkan"),
            ("Admin menekan tombol Login", "Sistem memproses autentikasi dan mengecek kecocokan data"),
            ("", "Jika kredensial benar, sistem mengarahkan admin ke halaman dashboard"),
            ("", "Sistem menampilkan semua elemen dashboard sesuai hak akses admin"),
        ],
        "alt_flow": [
            ("3a. Kredensial tidak valid",
             "Admin memasukkan username dan/atau kata sandi yang tidak sesuai. "
             "Sistem menampilkan pesan \"Username atau kata sandi salah\" dan proses login dihentikan."),
            ("3b. Field username atau kata sandi kosong",
             "Admin mengirimkan form login tanpa mengisi username dan/atau kata sandi. "
             "Sistem menampilkan pesan peringatan bahwa username dan kata sandi wajib diisi "
             "dan proses login dihentikan."),
        ],
        "extension": "-",
    },
    # ── UC-02 s/d UC-05 Mengelola Siswa ───────────────────────────
    {
        "id": "UC-02",
        "name": "Melihat Daftar Siswa",
        "description": (
            "Use case ini menggambarkan proses admin dalam menampilkan daftar siswa "
            "yang terdaftar beserta informasi profil dan capaian belajar pada modul "
            "Profil dan Kemajuan Siswa."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login ke sistem.",
        "post": "Daftar siswa tampil sesuai data terbaru.",
        "basic_flow": [
            ("Admin membuka menu Profil dan Kemajuan Siswa", "Sistem menampilkan halaman daftar siswa"),
            ("Admin membuka halaman Manajemen Daftar Siswa", "Sistem menampilkan tabel data siswa"),
            ("", "Sistem menampilkan nama, kelompok, total poin, dan bintang setiap siswa"),
        ],
        "alt_flow": [
            ("2a. Pencarian siswa",
             "Admin memasukkan kata kunci nama atau kelompok pada kolom pencarian. "
             "Sistem menampilkan daftar siswa yang sesuai dengan filter pencarian."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-03",
        "name": "Tambah Siswa",
        "description": (
            "Use case ini menggambarkan proses admin dalam menambahkan data siswa baru "
            "ke dalam sistem gamifikasi pembelajaran."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan berada pada halaman daftar siswa.",
        "post": "Data siswa baru berhasil disimpan dan ditampilkan pada daftar siswa.",
        "basic_flow": [
            ("Admin menekan tombol Tambah Siswa", "Sistem menampilkan form tambah data siswa"),
            ("Admin mengisi nama, kelompok, dan foto profil (opsional)", "Sistem memvalidasi kelengkapan data"),
            ("Admin menekan tombol Simpan", "Sistem menyimpan data siswa baru"),
            ("", "Sistem menampilkan daftar siswa terbaru"),
        ],
        "alt_flow": [
            ("3a. Data tidak lengkap",
             "Admin tidak mengisi field wajib. Sistem menampilkan pesan peringatan "
             "dan proses penyimpanan dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-04",
        "name": "Ubah Data Siswa",
        "description": (
            "Use case ini menggambarkan proses admin dalam memperbarui data profil siswa "
            "yang telah terdaftar pada sistem."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat data siswa pada sistem.",
        "post": "Data siswa berhasil diperbarui dan ditampilkan pada daftar siswa.",
        "basic_flow": [
            ("Admin memilih aksi edit pada siswa", "Sistem menampilkan form edit data siswa"),
            ("Admin mengubah nama, kelompok, atau foto profil", "Sistem memvalidasi data yang diubah"),
            ("Admin menekan tombol Simpan Perubahan", "Sistem memperbarui data siswa"),
            ("", "Sistem menampilkan daftar siswa terbaru"),
        ],
        "alt_flow": [
            ("3a. Data tidak valid",
             "Data yang diubah tidak lengkap atau tidak sesuai format. Sistem menampilkan "
             "pesan peringatan dan proses penyimpanan dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-05",
        "name": "Hapus Siswa",
        "description": (
            "Use case ini menggambarkan proses admin dalam menghapus data siswa "
            "dari sistem secara permanen."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat data siswa pada sistem.",
        "post": "Data siswa berhasil dihapus dan daftar siswa diperbarui.",
        "basic_flow": [
            ("Admin memilih aksi hapus pada siswa", "Sistem menampilkan dialog konfirmasi hapus"),
            ("Admin menekan tombol Hapus pada dialog konfirmasi", "Sistem menghapus data siswa"),
            ("", "Sistem menampilkan daftar siswa terbaru"),
        ],
        "alt_flow": [
            ("2a. Pembatalan hapus",
             "Admin menekan tombol Batal pada dialog konfirmasi. "
             "Sistem menutup dialog dan proses hapus dihentikan."),
        ],
        "extension": "-",
    },
    # ── UC-06 s/d UC-10 Mengelola Tema ────────────────────────────
    {
        "id": "UC-06",
        "name": "Melihat Daftar Tema",
        "description": (
            "Use case ini menggambarkan proses admin dalam menampilkan daftar tema "
            "pembelajaran beserta status aktivasi pada modul Manajemen Pembelajaran."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login ke sistem.",
        "post": "Daftar tema tampil sesuai data terbaru.",
        "basic_flow": [
            ("Admin membuka menu Manajemen Pembelajaran", "Sistem menampilkan menu modul pembelajaran"),
            ("Admin membuka halaman Manajemen Tema", "Sistem menampilkan daftar tema dalam bentuk kartu"),
            ("", "Sistem menampilkan nama, deskripsi, ikon, dan status aktif setiap tema"),
        ],
        "alt_flow": [
            ("2a. Pencarian tema",
             "Admin memasukkan kata kunci nama atau deskripsi tema. "
             "Sistem menampilkan tema yang sesuai dengan filter pencarian."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-07",
        "name": "Tambah Tema",
        "description": (
            "Use case ini menggambarkan proses admin dalam menambahkan tema pembelajaran "
            "baru ke dalam sistem."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan berada pada halaman manajemen tema.",
        "post": "Tema baru berhasil disimpan dan ditampilkan pada daftar tema.",
        "basic_flow": [
            ("Admin menekan tombol Tambah Tema", "Sistem menampilkan form tambah tema"),
            ("Admin mengisi nama tema, deskripsi, ikon, dan status tema", "Sistem memvalidasi kelengkapan data"),
            ("Admin menekan tombol Simpan", "Sistem menyimpan data tema baru"),
            ("", "Sistem menampilkan daftar tema terbaru"),
        ],
        "alt_flow": [
            ("3a. Data tidak lengkap",
             "Admin tidak mengisi field wajib. Sistem menampilkan pesan peringatan "
             "dan proses penyimpanan dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-08",
        "name": "Ubah Tema",
        "description": (
            "Use case ini menggambarkan proses admin dalam memperbarui informasi tema "
            "pembelajaran yang telah ada pada sistem."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat data tema pada sistem.",
        "post": "Data tema berhasil diperbarui dan ditampilkan pada daftar tema.",
        "basic_flow": [
            ("Admin memilih aksi edit pada tema", "Sistem menampilkan form edit tema"),
            ("Admin mengubah nama, deskripsi, ikon, atau status tema", "Sistem memvalidasi data yang diubah"),
            ("Admin menekan tombol Simpan", "Sistem memperbarui data tema"),
            ("", "Sistem menampilkan daftar tema terbaru"),
        ],
        "alt_flow": [
            ("3a. Data tidak valid",
             "Data yang diubah tidak lengkap. Sistem menampilkan pesan peringatan "
             "dan proses penyimpanan dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-09",
        "name": "Aktivasi Tema",
        "description": (
            "Use case ini menggambarkan proses admin dalam mengaktifkan atau "
            "menonaktifkan tema pembelajaran agar dapat atau tidak dapat digunakan "
            "dalam proses belajar siswa."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat data tema pada sistem.",
        "post": "Status aktivasi tema berhasil diperbarui.",
        "basic_flow": [
            ("Admin membuka halaman manajemen tema", "Sistem menampilkan daftar tema"),
            ("Admin mengubah toggle status Aktif/Nonaktif pada tema", "Sistem memperbarui status isActive tema"),
            ("", "Sistem menampilkan status tema terbaru pada kartu tema"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-10",
        "name": "Hapus Tema",
        "description": (
            "Use case ini menggambarkan proses admin dalam menghapus tema pembelajaran "
            "dari sistem secara permanen."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat data tema pada sistem.",
        "post": "Data tema berhasil dihapus dan daftar tema diperbarui.",
        "basic_flow": [
            ("Admin memilih aksi hapus pada tema", "Sistem menampilkan dialog konfirmasi hapus tema"),
            ("Admin menekan tombol Hapus pada dialog konfirmasi", "Sistem menghapus data tema"),
            ("", "Sistem menampilkan daftar tema terbaru"),
        ],
        "alt_flow": [
            ("2a. Pembatalan hapus",
             "Admin menekan tombol Batal. Sistem menutup dialog dan proses hapus dihentikan."),
        ],
        "extension": "-",
    },
    # ── UC-11 s/d UC-20 Mengelola Soal ────────────────────────────
    {
        "id": "UC-11",
        "name": "Melihat Daftar Soal Harian",
        "description": (
            "Use case ini menggambarkan proses admin dalam menampilkan daftar soal "
            "berdasarkan tema dan tanggal belajar pada modul Manajemen Soal."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah berhasil login dan terdapat minimal satu tema pembelajaran.",
        "post": "Daftar soal harian tampil sesuai tema dan tanggal yang dipilih.",
        "basic_flow": [
            ("Admin membuka halaman Manajemen Soal", "Sistem menampilkan daftar tema"),
            ("Admin memilih tema dan menekan Kelola Soal", "Sistem menampilkan kalender dan daftar tanggal belajar"),
            ("Admin memilih tanggal belajar", "Sistem menampilkan daftar soal pada tanggal tersebut"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-12",
        "name": "Tambah Soal",
        "description": (
            "Use case ini menggambarkan proses admin dalam menambahkan soal pembelajaran "
            "baru pada tema dan tanggal belajar tertentu."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman daftar soal harian tema dan tanggal tertentu.",
        "post": "Soal baru berhasil disimpan dan ditampilkan pada daftar soal.",
        "basic_flow": [
            ("Admin menekan tombol Tambah Soal", "Sistem menampilkan form tambah soal baru"),
            ("Admin memilih tipe soal dan mengisi instruksi, poin, serta media", "Sistem memvalidasi kelengkapan data soal"),
            ("Admin menekan tombol Tambah Soal", "Sistem menyimpan data soal baru"),
            ("", "Sistem menampilkan daftar soal terbaru"),
        ],
        "alt_flow": [
            ("3a. Data tidak lengkap",
             "Admin tidak mengisi field wajib. Sistem menampilkan pesan peringatan "
             "dan proses penyimpanan dihentikan."),
        ],
        "extension": "Kelola Opsi Jawaban (UC-18), Kelola Relasi Jawaban (UC-19), Kelola Kepingan Puzzle (UC-20)",
    },
    {
        "id": "UC-13",
        "name": "Ubah Soal",
        "description": (
            "Use case ini menggambarkan proses admin dalam memperbarui data soal "
            "pembelajaran yang telah ada pada sistem."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman daftar soal dan terdapat soal pada tanggal tersebut.",
        "post": "Data soal berhasil diperbarui dan ditampilkan pada daftar soal.",
        "basic_flow": [
            ("Admin memilih aksi edit pada soal", "Sistem menampilkan form edit soal"),
            ("Admin mengubah tipe, instruksi, poin, atau media soal", "Sistem memvalidasi data yang diubah"),
            ("Admin menekan tombol Simpan", "Sistem memperbarui data soal"),
            ("", "Sistem menampilkan daftar soal terbaru"),
        ],
        "alt_flow": [],
        "extension": "Kelola Opsi Jawaban (UC-18), Kelola Relasi Jawaban (UC-19), Kelola Kepingan Puzzle (UC-20)",
    },
    {
        "id": "UC-14",
        "name": "Aktifkan Soal Harian",
        "description": (
            "Use case ini menggambarkan proses admin dalam mengaktifkan ketersediaan "
            "soal harian agar dapat diakses oleh siswa pada tanggal belajar tertentu."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman kalender soal tema dan terdapat soal pada tanggal tersebut.",
        "post": "Status ketersediaan soal harian berhasil diperbarui.",
        "basic_flow": [
            ("Admin membuka halaman soal per tanggal", "Sistem menampilkan status ketersediaan soal"),
            ("Admin mengaktifkan ketersediaan soal pada tanggal", "Sistem memperbarui status isAvailable soal"),
            ("", "Sistem menampilkan status Aktif pada tanggal tersebut"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-15",
        "name": "Duplikat Soal",
        "description": (
            "Use case ini menggambarkan proses admin dalam menduplikasi seluruh soal "
            "dari satu tanggal ke tanggal belajar lainnya."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman kalender soal dan terdapat soal pada tanggal sumber.",
        "post": "Soal beserta opsi dan data terkait berhasil diduplikasi ke tanggal tujuan.",
        "basic_flow": [
            ("Admin memilih aksi duplikat pada tanggal sumber", "Sistem menampilkan form duplikat soal"),
            ("Admin memilih tanggal tujuan", "Sistem memvalidasi tanggal tujuan"),
            ("Admin menekan tombol Duplikat Soal", "Sistem menyalin soal beserta opsi dan data terkait"),
            ("", "Sistem menampilkan soal pada tanggal tujuan"),
        ],
        "alt_flow": [
            ("3a. Tanggal tujuan sudah memiliki soal",
             "Tanggal tujuan sudah berisi soal. Sistem menampilkan pesan peringatan "
             "dan proses duplikasi dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-16",
        "name": "Ubah Tanggal Soal",
        "description": (
            "Use case ini menggambarkan proses admin dalam memindahkan soal "
            "dari satu tanggal belajar ke tanggal belajar lainnya."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman kalender soal dan terdapat soal pada tanggal sumber.",
        "post": "Tanggal belajar soal berhasil diperbarui.",
        "basic_flow": [
            ("Admin memilih aksi ubah tanggal pada tanggal sumber", "Sistem menampilkan form ubah tanggal soal"),
            ("Admin memilih tanggal tujuan", "Sistem memvalidasi tanggal tujuan"),
            ("Admin menekan tombol Pindah Tanggal", "Sistem memperbarui learningDate seluruh soal"),
            ("", "Sistem menampilkan soal pada tanggal tujuan"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-17",
        "name": "Hapus Soal",
        "description": (
            "Use case ini menggambarkan proses admin dalam menghapus soal pembelajaran "
            "dari sistem secara permanen."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin berada pada halaman daftar soal dan terdapat soal pada tanggal tersebut.",
        "post": "Soal berhasil dihapus dan daftar soal diperbarui.",
        "basic_flow": [
            ("Admin memilih aksi hapus pada soal", "Sistem menampilkan dialog konfirmasi hapus"),
            ("Admin menekan tombol Hapus", "Sistem menghapus soal beserta opsi terkait"),
            ("", "Sistem menampilkan daftar soal terbaru"),
        ],
        "alt_flow": [
            ("2a. Pembatalan hapus",
             "Admin menekan tombol Batal. Sistem menutup dialog dan proses hapus dihentikan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-18",
        "name": "Kelola Opsi Jawaban",
        "description": (
            "Use case ini menggambarkan proses admin dalam mengelola opsi jawaban "
            "pada soal tipe pilihan ganda, mengurutkan, mencocokkan, dan seret-lepas."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah membuat soal dan berada pada halaman daftar soal.",
        "post": "Opsi jawaban soal berhasil dikelola.",
        "basic_flow": [
            ("Admin memilih aksi kelola opsi pada soal", "Sistem menampilkan form kelola opsi jawaban"),
            ("Admin menambah, mengubah, atau menghapus opsi jawaban", "Sistem memvalidasi data opsi"),
            ("Admin menandai jawaban benar atau urutan benar", "Sistem menyimpan perubahan opsi jawaban"),
            ("", "Sistem menampilkan opsi jawaban terbaru"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-19",
        "name": "Kelola Relasi Jawaban",
        "description": (
            "Use case ini menggambarkan proses admin dalam mengatur pasangan jawaban benar "
            "pada soal tipe mencocokkan dan seret-lepas."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah membuat soal tipe mencocokkan/seret-lepas dan opsi jawaban tersedia.",
        "post": "Pasangan jawaban benar berhasil disimpan.",
        "basic_flow": [
            ("Admin membuka kelola opsi soal tipe mencocokkan/seret-lepas", "Sistem menampilkan opsi pertanyaan dan jawaban"),
            ("Admin menentukan pasangan pertanyaan dan jawaban yang benar", "Sistem memvalidasi pasangan"),
            ("Admin menyimpan pasangan jawaban", "Sistem menyimpan data relasi matching"),
            ("", "Sistem menampilkan pasangan jawaban terbaru"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    {
        "id": "UC-20",
        "name": "Kelola Kepingan Puzzle",
        "description": (
            "Use case ini menggambarkan proses admin dalam mengatur konfigurasi puzzle "
            "dan kepingan gambar pada soal tipe puzzle."
        ),
        "actor": "Admin (Guru)",
        "pre": "Admin telah membuat soal tipe puzzle pada sistem.",
        "post": "Konfigurasi puzzle dan kepingan berhasil disimpan.",
        "basic_flow": [
            ("Admin membuka kelola puzzle pada soal", "Sistem menampilkan form konfigurasi puzzle"),
            ("Admin mengunggah gambar dan menentukan ukuran grid", "Sistem memvalidasi konfigurasi puzzle"),
            ("Admin menyimpan konfigurasi kepingan", "Sistem menyimpan data jigsaw puzzle dan kepingan"),
            ("", "Sistem menampilkan preview puzzle"),
        ],
        "alt_flow": [],
        "extension": "-",
    },
    # ── UC-21 s/d UC-25 Siswa ─────────────────────────────────────
    {
        "id": "UC-21",
        "name": "Pilih Profil Siswa",
        "description": (
            "Use case ini menggambarkan proses pemilihan profil siswa sebelum memulai "
            "aktivitas belajar pada sistem gamifikasi."
        ),
        "actor": "Siswa",
        "pre": "Siswa telah masuk ke halaman beranda sistem dan terdapat data siswa yang terdaftar.",
        "post": "Profil siswa terpilih dan sistem siap menampilkan menu pembelajaran.",
        "basic_flow": [
            ("Siswa menekan tombol mulai pada beranda", "Sistem menampilkan halaman daftar siswa"),
            ("Siswa mencari atau memilih profilnya", "Sistem menampilkan informasi profil siswa terpilih"),
            ("", "Sistem mengarahkan siswa ke halaman pilih tema"),
        ],
        "alt_flow": [
            ("2a. Profil siswa tidak ditemukan",
             "Siswa melakukan pencarian namun nama tidak ditemukan. Sistem menampilkan "
             "pesan bahwa siswa tidak ditemukan."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-22",
        "name": "Pilih Tema",
        "description": (
            "Use case ini menggambarkan proses pemilihan tema pembelajaran oleh siswa "
            "setelah profil dipilih, disertai informasi capaian belajar per tema."
        ),
        "actor": "Siswa",
        "pre": "Siswa telah memilih profil dan terdapat tema aktif pada sistem.",
        "post": "Tema pembelajaran terpilih dan sistem siap menampilkan peta belajar.",
        "basic_flow": [
            ("Siswa membuka halaman pilih tema", "Sistem menampilkan daftar tema aktif beserta skor dan bintang"),
            ("Siswa memilih tema dan menekan Lanjutkan", "Sistem menyimpan tema terpilih"),
            ("", "Sistem mengarahkan siswa ke halaman peta belajar"),
        ],
        "alt_flow": [
            ("2a. Tidak ada tema aktif",
             "Sistem tidak menemukan tema yang aktif. Sistem menampilkan pesan "
             "bahwa belum ada tema tersedia."),
        ],
        "extension": "-",
    },
    {
        "id": "UC-23",
        "name": "Navigasi Peta Belajar",
        "description": (
            "Use case ini menggambarkan proses navigasi siswa pada peta belajar untuk "
            "memilih tanggal belajar yang akan dikerjakan atau ditinjau."
        ),
        "actor": "Siswa",
        "pre": "Siswa telah memilih profil dan tema pembelajaran.",
        "post": "Siswa diarahkan ke halaman pengerjaan soal atau review sesuai status level.",
        "basic_flow": [
            ("Siswa membuka halaman peta belajar", "Sistem menampilkan level belajar per tanggal"),
            ("Siswa memilih level/tanggal belajar", "Sistem mengecek status level (belum/selesai dikerjakan)"),
            ("", "Jika belum selesai, sistem mengarahkan ke halaman kerjakan soal"),
            ("", "Jika sudah selesai, sistem mengarahkan ke halaman review jawaban"),
        ],
        "alt_flow": [
            ("2a. Tidak ada soal tersedia",
             "Tidak terdapat soal aktif. Sistem menampilkan pesan bahwa belum ada soal tersedia."),
        ],
        "extension": "Kerjakan Soal Harian (UC-24), Review Jawaban (UC-25)",
    },
    {
        "id": "UC-24",
        "name": "Kerjakan Soal Harian",
        "description": (
            "Use case ini menggambarkan proses pengerjaan soal harian oleh siswa, "
            "termasuk penilaian jawaban, pemberian poin, dan perhitungan capaian belajar."
        ),
        "actor": "Siswa",
        "pre": "Siswa telah memilih profil, tema, dan tanggal belajar yang memiliki soal tersedia.",
        "post": "Jawaban siswa tersimpan dan ringkasan hasil belajar ditampilkan.",
        "basic_flow": [
            ("Siswa membuka halaman soal harian", "Sistem menampilkan soal pertama sesuai tipe soal"),
            ("Siswa mengerjakan dan mengirim jawaban", "Sistem memvalidasi dan menilai jawaban"),
            ("", "Sistem menampilkan umpan balik benar/salah dan poin"),
            ("Siswa melanjutkan ke soal berikutnya", "Sistem menampilkan soal berikutnya hingga selesai"),
            ("", "Sistem menghitung total skor dan bintang"),
            ("", "Sistem menampilkan ringkasan hasil belajar"),
        ],
        "alt_flow": [
            ("4a. Waktu pengerjaan habis",
             "Soal memiliki batas waktu dan waktu habis. Sistem menghentikan pengerjaan "
             "dan melanjutkan ke proses penilaian."),
        ],
        "extension": "Pilihan Ganda, Mencocokkan, Mengurutkan, Seret dan Letakkan, Puzzle",
    },
    {
        "id": "UC-25",
        "name": "Review Jawaban",
        "description": (
            "Use case ini menggambarkan proses peninjauan kembali jawaban siswa setelah "
            "menyelesaikan soal harian, meliputi status benar/salah dan ringkasan capaian."
        ),
        "actor": "Siswa",
        "pre": "Siswa telah menyelesaikan soal harian pada tema dan tanggal tertentu.",
        "post": "Siswa dapat meninjau seluruh jawaban dan memahami capaian belajarnya.",
        "basic_flow": [
            ("Siswa membuka halaman review jawaban", "Sistem mengambil riwayat jawaban dan kunci jawaban"),
            ("", "Sistem menampilkan soal beserta status benar/salah"),
            ("Siswa menavigasi antar soal", "Sistem menampilkan detail jawaban dan poin per soal"),
            ("", "Sistem menampilkan ringkasan capaian belajar"),
        ],
        "alt_flow": [
            ("2a. Data jawaban tidak ditemukan",
             "Sistem tidak menemukan riwayat jawaban. Sistem menampilkan pesan "
             "bahwa data review belum tersedia."),
        ],
        "extension": "-",
    },
]


def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def format_alt_flow(alt_flows: list[tuple[str, str]]) -> str:
    if not alt_flows:
        return "-"
    parts = []
    for title, desc in alt_flows:
        parts.append(f"{title}\n{desc}")
    return "\n\n".join(parts)


def add_nested_basic_flow(cell, flows: list[tuple[str, str]]) -> None:
    """Nested 2-column table inside Basic Flow cell."""
    nested = cell.add_table(rows=1 + len(flows), cols=2)
    nested.style = "Table Grid"

    hdr = nested.rows[0].cells
    set_cell_text(hdr[0], "Actor's Action", bold=True)
    set_cell_text(hdr[1], "System's Response", bold=True)
    for c in hdr:
        set_cell_shading(c, "D9E2F3")

    step = 1
    for i, (actor, system) in enumerate(flows, start=1):
        row = nested.rows[i].cells
        left = f"{step}. {actor}" if actor.strip() else ""
        if actor.strip():
            step += 1
        right = f"{step}. {system}" if system.strip() else ""
        if system.strip():
            step += 1
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)


def add_use_case_table(doc: Document, uc: dict) -> None:
    """Satu tabel utuh per use case — format seperti contoh TA."""
    # 6 baris metadata + 1 basic flow + 1 alt + 1 extension = 9 baris
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"

    meta = [
        ("Use case ID Number", uc["id"]),
        ("Use case Name", uc["name"]),
        ("Brief Description", uc["description"]),
        ("Primary Actor", uc["actor"]),
        ("Pre-condition", uc["pre"]),
        ("Post Condition", uc["post"]),
    ]

    for i, (label, value) in enumerate(meta):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)

    # Basic Flow of Event — label kiri, sub-tabel kanan
    bf_label = table.rows[6].cells[0]
    bf_value = table.rows[6].cells[1]
    set_cell_text(bf_label, "Basic Flow of Event", bold=True)
    bf_value.text = ""
    add_nested_basic_flow(bf_value, uc["basic_flow"])

    # Alternative flow
    set_cell_text(table.rows[7].cells[0], "Alternative flow of events", bold=True)
    set_cell_text(table.rows[7].cells[1], format_alt_flow(uc["alt_flow"]))

    # Extension points
    set_cell_text(table.rows[8].cells[0], "Extension points", bold=True)
    set_cell_text(table.rows[8].cells[1], uc["extension"])

    # Lebar kolom label
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Spesifikasi Use Case\n")
    run.bold = True
    run.font.size = Pt(16)
    run2 = title.add_run("Sistem Gamifikasi Pembelajaran PAUD")
    run2.bold = True
    run2.font.size = Pt(14)
    doc.add_paragraph()

    for uc in USE_CASES:
        add_use_case_table(doc, uc)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
